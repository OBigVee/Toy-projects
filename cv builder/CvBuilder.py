from docx import Document
from docx.shared import Inches

class CvBuilder:
    
    def __init__(self):
        self.document = Document()
    
        
    def uploadPicture(self):
        """ method functions as picture uploader
        Returns
        -------
        None """
        self.document.add_picture("Me.png",width=Inches(1.0))
    
    def basicInfo(self):
        """
        function as a method of collecting basic info such as name etc...
        Returns
        -------
        None.
        """
        ## creates a paragraph for writing
        basic = self.document.add_paragraph()
        name = input("what is your name? :\n")
        ## format string 
        name = ("\t\t"+name.title()+"\t")
        ## write name out to file
        basic.add_run(name).bold=True
        
        while True:
            try:
                phoneNumber = int(input("enter your phone number:\n📞️:"))
                phoneNumber = ("📞️Tel:") + str(phoneNumber)+"\n"
                basic.add_run(phoneNumber).bold=True
            except ValueError as VE:
                print("phone number must be integer\n")
            else:
                break
        address = input("enter your location\n")
        address = ("\t\t\t🇳🇬️"+address.title()+"\n")
        basic.add_run(address).bold=True
                
        eMail = input("Enter your mail\n 📧️:")
        eMail = ("\t 📧️ Email:\t"+eMail+"\n")
        basic.add_run(eMail).bold=True
       
        gitHub = "www.github.com/"
        gitHubHandle = input("enter your github handle www.github.com/...")
        gitHub = ("\t\t" + gitHub + gitHubHandle+"\n")
        basic.add_run(gitHub).italic = True
        
        linkedIn = "www.linkedin.com/in/"
        linkedInHandle = input("enter your Linkedin handle www.linkedin.com/in/...")
        linkedIn = ("\t\t" + linkedIn + linkedInHandle+"\n")
        basic.add_run(linkedIn).italic = True 
        
        #self.document.add_paragraph(name + str(phoneNumber) + eMail + linkedIn + gitHub)
        
    def summary(self):
        """ this function takes input which users use to write a summary """
        self.document.add_heading("Summary")
        summary = input("enter a summary of yourself :\n")
        self.document.add_paragraph("\n"+summary)
        
    def experience(self):
        #self.document.add_heading("Work Experience")
        Experience = self.document.add_paragraph()
        company = input ("Enter compnay name? \n")
        startDate = input("start date? \n")
        endDate = input ("End date? \n")
        Experience.add_run(company + "\n").bold = True
        Experience.add_run(startDate+" - "+endDate+"\n").italic = True
        experienceDetails = input("describe your expirence at "+company+"\t\n:")
        Experience.add_run(experienceDetails+"\n")
        self.document.add_paragraph(self.hackingSkills())
        print("\n")
 
    def hackingSkills(self):
        """
        Funtion to add skills or tech stack used during a specifi work experience        
        Returns
        -------
        None.
        """
        Hacking = self.document.add_paragraph()
        self.document.add_paragraph()
        Hacking.add_run("hacking Skills 👨‍💻️: ").bold = True
        while True:
            skillQuest = input("do you want to add tech Stack? yes or no\n👨‍💻️")
            if skillQuest.lower() == 'yes':
                skill = input("enter hacking skill or technology stack\n👨‍💻️")
                Hacking.add_run(skill+"|")
            else:
                break
            print("\n")
            
    def projects(self,**description):
        proj = self.doument.add_paragraph()
        for key, value in description.items():
            proj.add_run(key).bold=True
            proj.add_run("\n"+value+"\n").italic=True
        
        return description
            
    def saveCV(self):
        """
        function to save file a with desired name
        Returns
        -------
        None."""
        askFileName = input("saved file as\n")+".docx"
        self.document.save(askFileName)
        
       
    ######################
       # # picuture
    # document.add_picture("Me.png",width=Inches(2.0))
    # # name deatails
    # name = input("enter your name::")
    # savedName = "\t"+name.title()+"\t"
    # phoneNumber = int(input("enter your phone number::"))
    # phoneNumber = "Tel:\t"+str(phoneNumber)+"\n"
    # mail = input("enter your Email::")
    # mail = "\tEmail:\t"+mail
    # document.add_paragraph(savedName+str(phoneNumber)+mail)
    # # Summary 
    # document.add_heading("Summary")
    # summary = input("enter a summary of yourself :")
    # document.add_paragraph("\n"+summary)
    # document.add_heading("Work Experience")
        #Experience = self.document.add_paragraph()
        #Hacking = self.document.add_paragraph()
        
    # More Experiences
    # createCV = CvBuilder()
    # createCV.uploadPicture()
    # createCV.basicInfo()
    # createCV.summary()
    # while True:
    #     hasMoreExperience = input("DO you want to add  experience ? \n Yes|NO")
    #     if hasMoreExperience.lower() == "yes":
    #         experience()
    #        # hackingSkills()
    #         print("\n")
    #     else:
    #         break
    # createCV.saveCV()
            
    
    #document.save("Custom CV.docx")
